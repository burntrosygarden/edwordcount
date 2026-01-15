#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
创建带有旧标注的测试文档
用于测试删除旧标注和排除标题字数的功能
"""

from docx import Document


def create_test_document_with_old_annotations():
    """创建包含旧标注的测试文档"""

    doc = Document()

    # 第一部分：引入（带旧标注）
    doc.add_paragraph("第一部分：引入(00:00 - 00:30)")
    doc.add_paragraph("""
大家好！今天我们要学习一个非常有趣的数学知识（播放开场动画）。
你们知道吗？数学其实就在我们的生活中【展示生活场景图片】。
让我们一起来探索吧！
""".strip())

    # 第二部分：知识点讲解（带旧标注）
    doc.add_paragraph("\n第二部分：知识点讲解（约500字，00:30 - 05:00）")

    doc.add_paragraph("知识点1：加法的基本概念")
    doc.add_paragraph("""
首先，我们来看看什么是加法（展示PPT第1页）。
加法就是把两个数字合在一起【动画演示：2个苹果+3个苹果】。
比如说，小明有2个苹果，小红给了他3个苹果(展示动画)，
那么小明现在一共有几个苹果呢？对了！是5个苹果「点赞动画」。
这就是2加3等于5，我们可以写成：2 + 3 = 5（板书展示）。
""".strip())

    doc.add_paragraph("\n知识点2：加法的运算方法")
    doc.add_paragraph("""
接下来，让我们学习如何计算加法［展示计算步骤］。
我们可以用手指来帮助计算｛演示手指计数｝。
比如计算 4 + 3，我们先伸出4根手指(动作演示)，
然后再伸出3根手指（慢动作展示），
数一数一共有几根？没错！是7根『鼓掌音效』。
""".strip())

    doc.add_paragraph("\n知识点3：练习与巩固")
    doc.add_paragraph("""
现在我们来做一些简单的练习题（展示练习题PPT）。
请小朋友们跟着老师一起算一算。
大家都算出来了吗「展示答案动画」？
""".strip())

    # 第三部分：综合练习（带旧标注）
    doc.add_paragraph("\n第三部分：综合练习（约300字，05:00 - 07:30）")

    doc.add_paragraph("练习题1")
    doc.add_paragraph("""
院子里有3只小猫(动画出现3只猫)，
又来了4只小猫（动画再出现4只猫），
现在一共有几只小猫呢？
""".strip())

    doc.add_paragraph("\n练习题2")
    doc.add_paragraph("""
摊上有6个橙子｛展示橙子图片｝，
老板又拿来2个橙子(动画效果)，
一共有多少个橙子？
""".strip())

    # 第四部分：总结（带旧标注）
    doc.add_paragraph("\n第四部分：总结(07:30-08:00)")
    doc.add_paragraph("""
今天我们学习了加法的知识（回顾PPT）。
我们知道了加法就是把数字合在一起【复习动画】。
希望小朋友们回家后多多练习「展示作业」！
下节课我们会学习减法，同样很有趣哦｛预告下节课内容｝。
""".strip())

    # 保存文档
    output_file = "测试视频脚本_带旧标注.docx"
    doc.save(output_file)
    print(f"✅ 测试文档已创建: {output_file}")

    # 打印说明
    print("\n测试文档特点:")
    print("1. 所有主标题都带有旧的时间标注（需要被删除并替换）")
    print("2. 包含子标题（知识点1、知识点2、练习题1等，不应计入字数）")
    print("3. 包含各种括号类型的后期提示（不应计入字数）")
    print("\n预期结果:")
    print("- 旧标注应该被删除")
    print("- 子标题行不计入字数")
    print("- 主标题行不计入字数")
    print("- 括号内容不计入字数")


if __name__ == '__main__':
    create_test_document_with_old_annotations()
