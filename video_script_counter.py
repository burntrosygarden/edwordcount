#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
视频脚本字数统计与时间预估工具
Video Script Counter

功能：
1. 读取Word文档（.docx）
2. 识别四个固定部分：引入、知识点讲解、综合练习、总结
3. 移除括号内容并统计字数
4. 基于220字/分钟计算时间轴
5. 在标题后添加字数和时间标注
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from copy import deepcopy


class VideoScriptCounter:
    """视频脚本字数统计与时间预估工具"""

    # 配置参数
    SPEECH_RATE = 220  # 儿童教学语速：220字/分钟

    # 支持的括号类型（所有常见的中英文括号）
    BRACKET_PAIRS = [
        ('(', ')'),
        ('（', '）'),
        ('[', ']'),
        ('【', '】'),
        ('「', '」'),
        ('『', '』'),
        ('{', '}'),
        ('｛', '｝'),
    ]

    # 四个固定部分的识别模式 - 每个部分可以有多个备选模式
    SECTION_PATTERNS = [
        # 第一部分：引入
        [
            r'第一部分[：:]\s*引入',
            r'^引入$',
            r'课程引入',
        ],
        # 第二部分：知识点讲解
        [
            r'第二部分[：:]\s*知识点讲解',
            r'^知识点讲解$',
            r'知识点',
        ],
        # 第三部分：综合练习
        [
            r'第三部分[：:]\s*综合练习',
            r'^综合练习$',
            r'练习',
        ],
        # 第四部分：总结
        [
            r'第四部分[：:]\s*总结',
            r'^总结$',
            r'课程总结',
            r'总结与结语',
            r'结语',
        ],
    ]

    def __init__(self, input_file):
        """初始化"""
        self.input_file = Path(input_file)
        if not self.input_file.exists():
            raise FileNotFoundError(f"文件不存在: {input_file}")

        if self.input_file.suffix.lower() != '.docx':
            raise ValueError("仅支持 .docx 格式文件")

        # 生成输出文件名
        output_name = self.input_file.stem + '_带标注.docx'
        self.output_file = self.input_file.parent / output_name

    def remove_brackets(self, text):
        """
        移除文本中所有括号及括号内的内容
        支持嵌套括号

        Args:
            text: 输入文本

        Returns:
            移除括号后的文本
        """
        result = text

        # 对每种括号类型进行处理
        for open_br, close_br in self.BRACKET_PAIRS:
            # 使用栈来处理嵌套括号
            while True:
                # 转义特殊字符
                open_escaped = re.escape(open_br)
                close_escaped = re.escape(close_br)

                # 查找最内层的括号对
                pattern = open_escaped + r'[^' + open_escaped + close_escaped + r']*?' + close_escaped
                new_result = re.sub(pattern, '', result)

                # 如果没有更多的括号，退出循环
                if new_result == result:
                    break
                result = new_result

        return result

    def remove_old_annotation(self, text):
        """
        删除标题中已有的时间标注（包括旧格式和新格式）

        匹配格式示例：
        - (00:45 - 06:30)
        - （00:45-06:30）
        - (约150字，00:45-06:30)
        - （约150字，00:45 - 06:30）
        - （约327字，00:00-01:29）
        - (157字)
        - (0:00 - 0:45)

        注意：工具会删除所有已有标注，然后添加最新计算的标注

        Args:
            text: 输入文本

        Returns:
            删除标注后的文本
        """
        # 匹配所有标注格式（包括旧格式和新格式），工具会添加最新计算的标注
        patterns = [
            # 匹配包含字数和时间的标注（新旧格式都删除）
            r'[（(]\s*约?\s*\d+\s*字\s*[,，]\s*\d{1,2}:\d{2}\s*[-\-~～]\s*\d{1,2}:\d{2}\s*[）)]',
            # 匹配只有时间范围的标注
            r'[（(]\s*\d{1,2}:\d{2}\s*[-\-~～]\s*\d{1,2}:\d{2}\s*[）)]',
            # 匹配只有字数的标注
            r'[（(]\s*约?\s*\d+\s*字\s*[）)]',
        ]

        result = text
        for pattern in patterns:
            result = re.sub(pattern, '', result)

        return result.strip()

    def is_subtitle(self, text):
        """
        判断是否是子标题（如"知识点1"、"知识点2"等）

        Args:
            text: 段落文本

        Returns:
            是否是子标题
        """
        # 匹配"知识点1"、"知识点2"、"练习题1"等格式
        subtitle_patterns = [
            r'^知识点\s*\d+',
            r'^练习题?\s*\d+',
            r'^例题\s*\d+',
            r'^第[一二三四五六七八九十\d]+题',
            r'^题目\s*\d+',
        ]

        text = text.strip()
        for pattern in subtitle_patterns:
            if re.match(pattern, text):
                return True
        return False

    def count_characters(self, text):
        """
        统计字符数（Word标准：中文字符数 + 英文单词数 + 数字）

        Args:
            text: 输入文本

        Returns:
            字符数
        """
        # 统计中文字符（包括汉字和中文标点）
        # Unicode范围：
        # \u4e00-\u9fff: CJK统一汉字
        # \u3000-\u303f: CJK符号和标点
        # \uff00-\uffef: 全角ASCII、全角标点
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', text))

        # 统计英文单词数（包括缩写词如Let's, don't等）
        # 匹配：字母 + 可选的撇号和字母
        english_words = len(re.findall(r"[a-zA-Z]+(?:'[a-zA-Z]+)?", text))

        # 统计数字
        digits = len(re.findall(r'\d', text))

        # 总字数 = 中文字符 + 英文单词 + 数字
        return chinese_chars + english_words + digits

    def format_time(self, seconds):
        """
        将秒数转换为时间格式 MM:SS 或 HH:MM:SS

        Args:
            seconds: 秒数（整数）

        Returns:
            格式化的时间字符串
        """
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        secs = seconds % 60

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"

    def calculate_duration(self, char_count):
        """
        根据字数计算时长（秒）

        Args:
            char_count: 字符数

        Returns:
            时长（秒，四舍五入）
        """
        # 字数 / (字/分钟) * 60 = 秒
        duration = (char_count / self.SPEECH_RATE) * 60
        return round(duration)  # 四舍五入到整秒

    def extract_section_text(self, doc, section_index):
        """
        提取指定部分的文本内容（排除标题行）

        Args:
            doc: Document对象
            section_index: 部分索引（0-3）

        Returns:
            (section_para_index, section_text): 部分起始段落索引和文本内容
        """
        section_patterns = self.SECTION_PATTERNS[section_index]  # 现在是列表
        next_section_patterns = self.SECTION_PATTERNS[section_index + 1] if section_index < 3 else None

        section_text = ""
        section_para_index = None
        in_section = False

        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            # 删除旧标注后再匹配
            cleaned_text = self.remove_old_annotation(para_text)

            # 检查是否是当前部分的开始 - 尝试所有备选模式
            if any(re.search(pattern, cleaned_text) for pattern in section_patterns):
                section_para_index = i
                in_section = True
                continue  # 跳过主标题行，不计入字数

            # 检查是否到达下一部分 - 尝试所有备选模式
            if in_section and next_section_patterns:
                if any(re.search(pattern, cleaned_text) for pattern in next_section_patterns):
                    break

            # 如果在当前部分内
            if in_section:
                # 跳过子标题行（如"知识点1"），不计入字数
                if self.is_subtitle(para_text):
                    continue

                # 添加正文内容
                section_text += para_text + "\n"

        return section_para_index, section_text

    def process_document(self):
        """处理文档，添加字数和时间标注"""
        print(f"正在处理文件: {self.input_file}")

        # 读取文档
        doc = Document(str(self.input_file))

        # 存储每个部分的统计信息
        sections_info = []
        cumulative_time = 0  # 累积时间（秒）

        # 处理四个部分
        section_names = ['引入', '知识点讲解', '综合练习', '总结']

        for i, section_name in enumerate(section_names):
            print(f"\n处理第{i+1}部分：{section_name}")

            # 提取部分文本
            para_index, section_text = self.extract_section_text(doc, i)

            if para_index is None:
                print(f"  ⚠️  未找到该部分")
                continue

            # 移除括号内容
            text_without_brackets = self.remove_brackets(section_text)

            # 统计字数
            char_count = self.count_characters(text_without_brackets)

            # 计算时长
            duration = self.calculate_duration(char_count)

            # 计算时间范围
            start_time = cumulative_time
            end_time = cumulative_time + duration
            cumulative_time = end_time

            # 格式化时间
            start_str = self.format_time(start_time)
            end_str = self.format_time(end_time)

            # 保存信息
            sections_info.append({
                'index': i,
                'name': section_name,
                'para_index': para_index,
                'char_count': char_count,
                'duration': duration,
                'time_range': f"{start_str}-{end_str}"
            })

            print(f"  ✓ 字数: {char_count}")
            print(f"  ✓ 时长: {duration}秒")
            print(f"  ✓ 时间轴: {start_str}-{end_str}")

        # 在文档中添加标注
        print("\n\n正在生成带标注的文档...")
        for info in sections_info:
            para = doc.paragraphs[info['para_index']]

            # 删除旧的时间标注
            clean_text = self.remove_old_annotation(para.text)

            # 构建新的标注文本
            annotation = f"（约{info['char_count']}字，{info['time_range']}）"

            # 添加新标注
            para.text = clean_text + annotation

        # 保存文档
        doc.save(str(self.output_file))
        print(f"\n✅ 处理完成！")
        print(f"输出文件: {self.output_file}")

        # 打印总结
        print("\n" + "="*50)
        print("统计摘要")
        print("="*50)
        total_chars = sum(info['char_count'] for info in sections_info)
        total_duration = sum(info['duration'] for info in sections_info)
        print(f"总字数: {total_chars} 字")
        print(f"总时长: {self.format_time(total_duration)} ({total_duration}秒)")
        print(f"平均语速: {self.SPEECH_RATE} 字/分钟")
        print("="*50)


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("使用方法: python video_script_counter.py <输入文件.docx>")
        print("\n示例:")
        print("  python video_script_counter.py 我的视频脚本.docx")
        sys.exit(1)

    input_file = sys.argv[1]

    try:
        counter = VideoScriptCounter(input_file)
        counter.process_document()
    except Exception as e:
        print(f"\n❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
